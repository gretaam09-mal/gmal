import io

from docx import Document

from services.exports.docx import render_memo_docx

_CONTENT = {
    "headline": {"low": "20000.00", "likely": "25000.00", "high": "32500.00", "currency": "GBP"},
    "confidence_grade": "B",
    "confidence_score": "78.5",
    "headline_summary": "This deal carries a bounded, quantified regulatory exposure.",
    "obligations": [
        {
            "obligation_summary": "Appoint a data protection officer.",
            "clause_refs": ["s.1"],
            "impact_low": "20000.00",
            "impact_likely": "25000.00",
            "impact_high": "32500.00",
            "currency": "GBP",
            "present_value": "24000.00",
            "what_it_requires": "Appoint a data protection officer.",
            "why_it_applies": "The target processes personal data at scale.",
        }
    ],
    "excluded": [
        {
            "obligation_summary": "Client money segregation.",
            "rationale": "Does not bind: does not hold client money.",
        }
    ],
    "excluded_summary": "One obligation was excluded — see below.",
}


def _paragraphs(docx_bytes: bytes) -> list[str]:
    document = Document(io.BytesIO(docx_bytes))
    return [p.text for p in document.paragraphs]


def test_render_memo_docx_produces_a_readable_document():
    docx_bytes = render_memo_docx(memo_title="Project Falcon — Impact Memo", content=_CONTENT)

    document = Document(io.BytesIO(docx_bytes))
    headings = [p.text for p in document.paragraphs if p.style.name.startswith("Heading")]
    assert "Project Falcon — Impact Memo" in headings


def test_render_memo_docx_includes_headline_and_confidence():
    docx_bytes = render_memo_docx(memo_title="Project Falcon", content=_CONTENT)
    text = " ".join(_paragraphs(docx_bytes))
    assert "25000.00" in text
    assert "Confidence grade: B" in text


def test_render_memo_docx_includes_obligations_and_exclusions():
    docx_bytes = render_memo_docx(memo_title="Project Falcon", content=_CONTENT)
    text = " ".join(_paragraphs(docx_bytes))
    assert "Appoint a data protection officer." in text
    assert "s.1" in text
    assert "Client money segregation." in text


def test_render_memo_docx_is_deterministic():
    first = render_memo_docx(memo_title="Project Falcon", content=_CONTENT)
    second = render_memo_docx(memo_title="Project Falcon", content=_CONTENT)
    assert _paragraphs(first) == _paragraphs(second)
