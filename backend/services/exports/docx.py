"""F8: renders a memo version's body — headline, obligations, exclusions
— as a DOCX for dropping into an IC pack. Deliberately body-only (no
assumption register or lineage appendix, which are PDF-only per F8's
spec); pure, takes only already-fetched content, no I/O beyond the
in-memory buffer it returns.
"""
from __future__ import annotations

import io

from docx import Document


def render_memo_docx(*, memo_title: str, content: dict) -> bytes:
    document = Document()
    document.add_heading(memo_title, level=1)
    document.add_paragraph(
        f"Confidence grade: {content['confidence_grade']} ({content['confidence_score']})"
    )

    headline = content["headline"]
    document.add_heading("Headline exposure", level=2)
    document.add_paragraph(
        f"Best {headline['low']} {headline['currency']} · "
        f"Likely {headline['likely']} {headline['currency']} · "
        f"Worst {headline['high']} {headline['currency']}"
    )
    document.add_paragraph(content["headline_summary"])

    document.add_heading("Obligations", level=2)
    for obligation in content["obligations"]:
        document.add_heading(obligation["obligation_summary"], level=3)
        document.add_paragraph(obligation["what_it_requires"])
        document.add_paragraph(obligation["why_it_applies"])
        document.add_paragraph("Clauses: " + ", ".join(obligation["clause_refs"]))
        document.add_paragraph(
            f"Cost: best {obligation['impact_low']}, likely {obligation['impact_likely']}, "
            f"worst {obligation['impact_high']} {obligation['currency']}"
        )
        document.add_paragraph(f"Present value: {obligation['present_value']}")

    document.add_heading("Excluded obligations", level=2)
    document.add_paragraph(content["excluded_summary"])
    for item in content["excluded"]:
        document.add_paragraph(f"{item['obligation_summary']} — {item['rationale']}")

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
